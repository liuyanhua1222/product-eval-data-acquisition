"""
generate_report.py — 生成产品评估报告

鉴权模式: nologin
依赖: python-docx（仅 docx 格式需要，pip install python-docx）

用法:
  python generate_report.py --result-file .cms-log/log/.../evaluation-xxx.json
  python generate_report.py --result-file evaluation.json --format docx --output report.docx
  python generate_report.py --result-file evaluation.json --format markdown
"""

import argparse
import json
import sys
from datetime import datetime
from pathlib import Path

LOG_DIR = Path(".cms-log/log/product-eval-data-acquisition")
LOG_DIR.mkdir(parents=True, exist_ok=True)


def result_icon(result: str) -> str:
    """将评估结论转换为图标。"""
    icons = {
        "通过": "✅",
        "不通过": "❌",
        "需补充数据": "⚠️",
        "待执行": "⏳",
        "待采集或判读": "⏳",
        "待人工判读": "🧠",
        "不适用": "—",
    }
    return icons.get(result, "❓")


def generate_markdown(data: dict) -> str:
    """生成 Markdown 格式评估报告。"""
    product = data.get("product", "未知产品")
    eval_date = data.get("evaluation_date", datetime.now().strftime("%Y-%m-%d"))
    verdict = data.get("final_verdict", "待完成")
    terminated_by = data.get("terminated_by")

    lines = [
        f"# {product} 新品评估报告",
        f"",
        f"**框架版本**: 3.1 | **评估日期**: {eval_date} | **最终结论**: {verdict}",
        f"",
    ]

    if terminated_by:
        lines += [
            f"> 🛑 **评估已终止**: 触发硬性 NO GO 规则 `{terminated_by}`",
            f"",
        ]

    for stage_key, stage_data in data.get("stage_results", {}).items():
        stage_num = stage_key.replace("stage_", "")
        stage_names = {"1": "第一阶段：入门筛选", "2": "第二阶段：深度评估", "3": "第三阶段：最终放行"}
        stage_title = stage_names.get(stage_num, f"第 {stage_num} 阶段")

        lines += [
            f"## {stage_title}",
            f"",
            f"| 规则 | 名称 | 适用性 | 结论 | 关键数值 | 证据 | 说明 | 缺失来源 |",
            f"|------|------|--------|------|---------|------|------|---------|",
        ]

        for item in stage_data.get("items", []):
            applicable = "适用" if item.get("applicable") else ("不适用" if item.get("applicable") is False else "—")
            result = item.get("result", "待执行")
            icon = result_icon(result)
            key_value = item.get("key_value") or "—"
            evidence = item.get("evidence") or "—"
            note = item.get("note") or "—"
            missing = item.get("missing_source") or "无"

            lines.append(
                f"| `{item['rule']}` | {item['name']} | {applicable} | {icon} {result} "
                f"| {key_value} | {evidence} | {note} | {missing} |"
            )

        lines.append("")

    lines += [
        "---",
        f"*报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}*",
    ]

    return "\n".join(lines)


def generate_docx(data: dict, output_path: Path) -> None:
    """生成 Word 格式评估报告。"""
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("❌ 缺少依赖: pip install python-docx", file=sys.stderr)
        sys.exit(1)

    product = data.get("product", "未知产品")
    eval_date = data.get("evaluation_date", datetime.now().strftime("%Y-%m-%d"))
    verdict = data.get("final_verdict", "待完成")

    doc = Document()

    # 标题
    title = doc.add_heading(f"{product} 新品评估报告", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 基本信息
    doc.add_paragraph(f"框架版本: 3.1 | 评估日期: {eval_date} | 最终结论: {verdict}")

    terminated_by = data.get("terminated_by")
    if terminated_by:
        p = doc.add_paragraph(f"⚠️ 评估已终止：触发硬性 NO GO 规则 {terminated_by}")
        p.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    # 各阶段结果
    stage_names = {"1": "第一阶段：入门筛选", "2": "第二阶段：深度评估", "3": "第三阶段：最终放行"}
    for stage_key, stage_data in data.get("stage_results", {}).items():
        stage_num = stage_key.replace("stage_", "")
        doc.add_heading(stage_names.get(stage_num, f"第 {stage_num} 阶段"), level=2)

        items = stage_data.get("items", [])
        if not items:
            doc.add_paragraph("（无评估事项）")
            continue

        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["规则", "名称", "结论", "关键数值", "证据", "缺失来源"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h

        for item in items:
            row = table.add_row()
            row.cells[0].text = item.get("rule", "")
            row.cells[1].text = item.get("name", "")
            row.cells[2].text = item.get("result", "待执行")
            row.cells[3].text = item.get("key_value") or "—"
            row.cells[4].text = item.get("evidence") or "—"
            row.cells[5].text = item.get("missing_source") or "无"

    doc.add_paragraph(f"\n报告生成时间: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.save(str(output_path))
    print(f"✅ Word 报告已保存: {output_path}", file=sys.stderr)


def main() -> None:
    parser = argparse.ArgumentParser(description="生成产品评估报告")
    parser.add_argument("--result-file", required=True, help="run_evaluation.py 输出的 JSON 文件路径")
    parser.add_argument(
        "--format",
        choices=["markdown", "docx"],
        default="markdown",
        help="输出格式：markdown（默认）或 docx",
    )
    parser.add_argument("--output", help="输出文件路径（不传则输出到 stdout）")
    args = parser.parse_args()

    result_path = Path(args.result_file)
    if not result_path.exists():
        print(f"❌ 文件不存在: {args.result_file}", file=sys.stderr)
        sys.exit(1)

    data = json.loads(result_path.read_text(encoding="utf-8"))
    product = data.get("product", "report")

    if args.format == "docx":
        output_path = Path(args.output) if args.output else LOG_DIR / f"{product}-report.docx"
        generate_docx(data, output_path)
    else:
        content = generate_markdown(data)
        if args.output:
            Path(args.output).write_text(content, encoding="utf-8")
            print(f"✅ Markdown 报告已保存: {args.output}", file=sys.stderr)
        else:
            print(content)


if __name__ == "__main__":
    main()
